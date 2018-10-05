# -*- coding: utf-8 -*-
from __future__ import unicode_literals

from django.shortcuts import render, get_object_or_404, redirect
from django.http import HttpResponse
from django.template import loader
from sertifikasiguru.models import Guru, File, DetailOrder, Order, Paket
from sertifikasiguru.decorators import decoclient
from sertifikasiguru.forms import LoginClientForm
from sertifikasiguru.MyCustomBackend import MyCustomBackend
from django.contrib import messages
from django.views.decorators.csrf import csrf_exempt
from django.http import HttpResponseRedirect
from django.contrib import messages
from sertifikasiguru import settings

def custom_redirect(url_name, **kwargs):
    from django.urls import reverse 
    import urllib
    url = reverse(url_name)
    params = urllib.parse.urlencode(kwargs)
    return HttpResponseRedirect(url + "?%s" % params)


def replacing (request):
	instance = get_object_or_404(DetailOrder, id_detailorder=request.GET['id'])
	messages.success(request, 'Generate detailOrder '+str(instance.id_detailorder))
	return replaceyaya(instance)

def replaceyaya(detailOrder):
	print(detailOrder)
	print(detailOrder.file.file.path)
	f = open(detailOrder.file.file.path , 'rb')
	document = Document(f)
	f.close()
	# xnmk	: Nama Kepala
	# xnmg	: Nama Guru
	# xnik	: NIP Kepala
	# xnig	: NIP Guru
	# xskn	: Nama Sekolah
	# xska	: ALamat Sekolah
	# xtdk	: TTD Kepala
	# xtdg	: TTD Guru
	# unmk	: Nama Kepala with underline
	# unmg	: Nama Guru with underline
	filename= detailOrder.file.filename
	nama_guru= detailOrder.order.guru.guru
	nama_kepala= detailOrder.order.kepala.kepala
	nip_guru= detailOrder.order.guru.nip_guru
	nip_kepala= detailOrder.order.kepala.nip_kepala
	nama_sekolah= detailOrder.order.sekolah.sekolah
	alamat_sekolah= detailOrder.order.sekolah.alamat
	ttd_kepala= detailOrder.order.kepala.ttd_kepala.url.strip("/")
	ttd_guru= detailOrder.order.guru.ttd_guru.url.strip("/")

	for p in document.paragraphs:
		if 'xnmk' in p.text or 'xnmg' in p.text or 'xnik' in p.text or 'xnig' in p.text or 'xskn' in p.text or 'xska' in p.text or 'xtdk' in p.text or 'xtdg' in p.text or 'unmk' in p.text or 'unmg' in p.text:
			inline = p.runs
			for i in range(len(inline)):
				if 'unmg' in inline[i].text:
					text = inline[i].text.replace('unmg', nama_guru)
					inline[i].text = text
					inline[i].underline = True
				if 'unmk' in inline[i].text:
					text = inline[i].text.replace('unmk', nama_kepala)
					inline[i].text = text
					inline[i].underline = True
				if 'xnik' in inline[i].text:
					text = inline[i].text.replace('xnik', nip_kepala)
					inline[i].text = text
				if 'xnig' in inline[i].text:
					text = inline[i].text.replace('xnig', nip_guru)
					inline[i].text = text
				if 'xnmk' in inline[i].text:
					text = inline[i].text.replace('xnmk', nama_kepala)
					inline[i].text = text
				if 'xnmg' in inline[i].text:
					text = inline[i].text.replace('xnmg', nama_guru)
					inline[i].text = text
				if 'xskn' in inline[i].text:
					text = inline[i].text.replace('xskn', nama_sekolah)
					inline[i].text = text
				if 'xska' in inline[i].text:
					text = inline[i].text.replace('xska', alamat_sekolah)
					inline[i].text = text
				if 'xtdk' in inline[i].text:
					inline[i].text = ''
					inline[i].add_picture(ttd_kepala)
				if 'xtdg' in inline[i].text:
					inline[i].text = ''
					inline[i].add_picture(ttd_guru)

	for table in document.tables:
		for row in table.rows:
			for cell in row.cells:
				for p in cell.paragraphs:
					if 'xnmk' in p.text or 'xnmg' in p.text or 'xnik' in p.text or 'xnig' in p.text or 'xskn' in p.text or 'xska' in p.text or 'xtdk' in p.text or 'xtdg' in p.text or 'unmk' in p.text or 'unmg' in p.text:
						inline = p.runs
						for i in range(len(inline)):
							if 'unmg' in inline[i].text:
								text = inline[i].text.replace('unmg', nama_guru)
								inline[i].text = text
								inline[i].underline = True
							if 'unmk' in inline[i].text:
								text = inline[i].text.replace('unmk', nama_kepala)
								inline[i].text = text
								inline[i].underline = True
							if 'xnik' in inline[i].text:
								text = inline[i].text.replace('xnik', nip_kepala)
								inline[i].text = text
							if 'xnig' in inline[i].text:
								text = inline[i].text.replace('xnig', nip_guru)
								inline[i].text = text
							if 'xnmk' in inline[i].text:
								text = inline[i].text.replace('xnmk', nama_kepala)
								inline[i].text = text
							if 'xnmg' in inline[i].text:
								text = inline[i].text.replace('xnmg', nama_guru)
								inline[i].text = text
							if 'xskn' in inline[i].text:
								text = inline[i].text.replace('xskn', nama_sekolah)
								inline[i].text = text
							if 'xska' in inline[i].text:
								text = inline[i].text.replace('xska', alamat_sekolah)
								inline[i].text = text
							if 'xtdk' in inline[i].text:
								inline[i].text = ''
								inline[i].add_picture(ttd_kepala)
							if 'xtdg' in inline[i].text:
								inline[i].text = ''
								inline[i].add_picture(ttd_guru)

	# response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
	# response['Content-Disposition'] = 'attachment; filename=download.docx'
	# document.save(response)
	document.save('media/processed/'+str(detailOrder.id_detailorder)+'.docx')
	detailOrder.processed= settings.MEDIA_URL+'processed/'+str(detailOrder.id_detailorder)+'.docx'

	detailOrder.save()
	
	# return response
	return custom_redirect('detailorder', id = detailOrder.order.id_order)

@csrf_exempt
def testdocx(request):
	template = loader.get_template('testdocx.html')
	context = {
		'page': 'Informasi',
		'data': 'Main',
	}

	if request.method == 'POST':
		f = open('media/test.docx', 'rb')
		document = Document(f)
		f.close()
		for p in document.paragraphs:
			if 'markng' in p.text or 'markig' in p.text or 'marknk' in p.text or 'markik' in p.text or 'marks' in p.text or 'marka' in p.text:
				inline = p.runs
				for i in range(len(inline)):
					if 'markng' in inline[i].text:
						text = inline[i].text.replace('markng', request.POST['nama_guru'])
						inline[i].text = text
					if 'markig' in inline[i].text:
						text = inline[i].text.replace('markig', request.POST['nip_guru'])
						inline[i].text = text
					if 'marknk' in inline[i].text:
						text = inline[i].text.replace('marknk', request.POST['nama_kepala'])
						inline[i].text = text
					if 'markik' in inline[i].text:
						text = inline[i].text.replace('markik', request.POST['nip_kepala'])
						inline[i].text = text
					if 'marks' in inline[i].text:
						text = inline[i].text.replace('marks', request.POST['sekolah'])
						inline[i].text = text
					if 'marka' in inline[i].text:
						text = inline[i].text.replace('marka', request.POST['alamat'])
						inline[i].text = text

		for table in document.tables:
			for row in table.rows:
				for cell in row.cells:
					for p in cell.paragraphs:
						if 'markng' in p.text or 'markig' in p.text or 'marknk' in p.text or 'markik' in p.text or 'marks' in p.text or 'marka' in p.text or 'ttdg' in p.text or 'ttdk' in p.text or 'markung' in p.text or 'markunk' in p.text:
							inline = p.runs
							for i in range(len(inline)):
								if 'markung' in inline[i].text:
									text = inline[i].text.replace('markung', request.POST['nama_guru'])
									inline[i].text = text
									inline[i].underline = True
								if 'markunk' in inline[i].text:
									text = inline[i].text.replace('markunk', request.POST['nama_kepala'])
									inline[i].text = text
									inline[i].underline = True
								if 'markng' in inline[i].text:
									text = inline[i].text.replace('markng', request.POST['nama_guru'])
									inline[i].text = text
								if 'markig' in inline[i].text:
									text = inline[i].text.replace('markig', request.POST['nip_guru'])
									inline[i].text = text
								if 'marknk' in inline[i].text:
									text = inline[i].text.replace('marknk', request.POST['nama_kepala'])
									inline[i].text = text
								if 'markik' in inline[i].text:
									text = inline[i].text.replace('markik', request.POST['nip_kepala'])
									inline[i].text = text
								if 'marks' in inline[i].text:
									text = inline[i].text.replace('marks', request.POST['sekolah'])
									inline[i].text = text
								if 'marka' in inline[i].text:
									text = inline[i].text.replace('marka', request.POST['alamat'])
									inline[i].text = text
								if 'ttdg' in inline[i].text:
									inline[i].text = ''
									inline[i].add_picture('media/ttd/ijazah_001.jpg')
								if 'ttdk' in inline[i].text:
									inline[i].text = ''
									inline[i].add_picture('media/ttd/ijazah_001.jpg')
								
		# paragraph = document.add_paragraph('Lorem ipsum dolor sit amet.')
		# run = paragraph.add_run()
		# # inline_shape = run.add_inline_picture('media/ttd/23.gif', MIME_type=None)
		# run.add_picture('media/ttd/23.gif')

		response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
		response['Content-Disposition'] = 'attachment; filename=download.docx'
		document.save(response)

		return response

	return HttpResponse(template.render(context, request))

from docx import Document
from docx.shared import Inches

def downloaddocx(request):
	document = Document()
	document.add_heading('Document Title', 0)

	response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
	response['Content-Disposition'] = 'attachment; filename=download.docx'
	document.save(response)

	return response

def readdocx(request):
	document = Document()
	document.add_heading('Document Title', 0)

	response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
	response['Content-Disposition'] = 'attachment; filename=download.docx'
	document.save(response)

	return response

def readwritedocx(request):
	document = Document()
	document.add_heading('Document Title', 0)

	response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
	response['Content-Disposition'] = 'attachment; filename=download.docx'
	document.save(response)

	return response

@decoclient
def dokumen(request):
	template = loader.get_template('client/v_dokumen.html')
	context = {
		'page': 'Dokumen',
		'client': Guru.objects.get(id_guru=request.session['client']),
    }
	return HttpResponse(template.render(context, request))

def login(request):
	template = loader.get_template('client/cl_login.html')
	if request.method == 'POST':
		form = LoginClientForm(request.POST)
		if form.is_valid():
			# form.save()
			mcb=MyCustomBackend()
			# if  mcb.authenticateadmin(form.cleaned_data.get('username'),form.cleaned_data.get('password')):
			client = mcb.authenticateclient(form.cleaned_data.get('email'),form.cleaned_data.get('password'))
			# client = authenticate(username=form.cleaned_data.get('username'), password=form.cleaned_data.get('password'))
			if  client:
				request.session['client'] = client.id_guru
				return redirect('dokumen')
			else:
				messages.error(request, 'Tidak bisa masuk')
				return redirect('masuk')
	else:
		form = LoginClientForm()
	
	# return render(request, 'login.html', {'form': form})
	return HttpResponse(template.render({'form': form}, request))

@decoclient
def logout(request):
	request.session.flush()
	return redirect('dokumen')